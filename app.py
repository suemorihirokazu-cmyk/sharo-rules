import streamlit as st
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import io
import re
import pandas as pd

# ==========================================================
# --- 全角/半角 変換用テーブル ---
# ==========================================================
HALF_CHARS = "0123456789abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ!\"#$%&'()*+,-./:;<=>?@[\\]^_`{|}~¥"
FULL_CHARS = "０１２３４５６７８９ａｂｃｄｅｆｇｈｉｊｋｌｍｎｏｐｑｒｓｔｕｖｗｘｙｚＡＢＣＤＥＦＧＨＩＪＫＬＭＮＯＰＱＲＳＴＵＶＷＸＹＺ！＂＃＄％＆’（）＊＋，－．／：；＜＝＞？＠［＼］＾＿｀｛｜｝～￥"

HALF2FULL = str.maketrans(HALF_CHARS, FULL_CHARS)
FULL2HALF = str.maketrans(FULL_CHARS, HALF_CHARS)

BORDER_CHARS = ['┌', '└', '│', '─', '├', '┤', '┬', '┴', '┼', '━', '┃', '┏', '┓', '┗', '┛', '┣', '┫', '┳', '┻', '╋', '┠', '┨', '┯', '┷', '┿', '┝', '┥', '┰', '┸', '╂']

# --- 階層ごとの正規表現パターン ---
RE_SHOU = r'^第[ \t　]*[0-9１２３４５６７８９０一二三四五六七八九十百]+[ \t　]*章'
RE_JOU  = r'^第[ \t　]*[0-9１２３４５６７８９０一二三四五六七八九十百]+[ \t　]*条'
# 【修正箇所】数字＋ピリオド、または、数字＋スペース（全角・半角・タブ）の場合に項として判定
RE_KOU  = r'^[０-９0-9]+(?:[．\.]|[ \t　]+)'
RE_GOU  = r'^[①-⑳㉑-㉟㊱-㊿]'
RE_LV5  = r'^[（\(][０-９0-9]+[）\)]'
RE_LV6  = r'^[（\(][ア-ンァ-ォ]+[）\)]'
RE_MIDASHI_GO = r'^[（\(].+[）\)]$' # カッコで囲まれた見出し語

TITLE_LIST = ["就　業　規　規則", "就業規則", "賃　金　規　程", "賃金規程", "育児・介護休業規程", "退職金規程", "付　則", "付則"]

def convert_char_width(text, mode):
    if mode == "半角に統一":
        return text.translate(FULL2HALF)
    elif mode == "全角に統一（推奨）":
        return text.translate(HALF2FULL)
    return text

def get_hierarchy_match(line, regex_pat, custom_markers_str, width_mode):
    m = re.match(regex_pat, line)
    if m:
        return m.group(0)
    if custom_markers_str:
        converted_custom = convert_char_width(custom_markers_str, width_mode)
        markers = [m.strip() for m in converted_custom.split(',') if m.strip()]
        for marker in markers:
            if line.startswith(marker):
                return marker
    return None

def strip_marker(text, marker_str):
    if not marker_str: return text
    if text.startswith(marker_str):
        return text[len(marker_str):].lstrip()
    return text

def normalize_marker_space(line, marker):
    if not marker: return line
    body = line[len(marker):].lstrip(" \t　")
    if body:
        return f"{marker}{body}"
    return marker

def renumber_headings(text):
    chapter_count = 0
    article_count = 0
    new_lines = []
    change_log = []
    
    pat_chapter = re.compile(r'^([ \t　]*)第[ \t　]*[0-9１２３４５６７８９０一二三四五六七八九十百]+[ \t　]*章(.*)')
    pat_article = re.compile(r'^([ \t　]*)第[ \t　]*[0-9１２３４５６７８９０一二三四五六七八九十百]+[ \t　]*条(の[0-9１２３４５６７８９０一二三四五六七八九十百]+)?(.*)')

    for line in text.split('\n'):
        original_line = line
        m_chap = pat_chapter.match(line)
        m_art = pat_article.match(line)
        
        if m_chap:
            chapter_count += 1
            new_line = f"{m_chap.group(1)}第{chapter_count}章{m_chap.group(2)}"
            if original_line.strip() != new_line.strip():
                change_log.append({"階層": "章", "変更前 (旧)": original_line.strip(), "変更後 (新)": new_line.strip()})
            line = new_line
            
        elif m_art:
            if not m_art.group(2):
                article_count += 1
            branch = m_art.group(2) if m_art.group(2) else ""
            new_line = f"{m_art.group(1)}第{article_count}条{branch}{m_art.group(3)}"
            
            if original_line.strip() != new_line.strip():
                change_log.append({"階層": "条", "変更前 (旧)": original_line.strip(), "変更後 (新)": new_line.strip()})
            line = new_line
        
        new_lines.append(line)
        
    return '\n'.join(new_lines), change_log

def set_font(run, font_name, size_pt=None, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if size_pt:
        run.font.size = Pt(size_pt)
    if bold:
        run.bold = True

def add_toc(doc):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.snap_to_grid = False
    run = paragraph.add_run("目　次")
    set_font(run, 'MS ゴシック', size_pt=14, bold=True)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.paragraph_format.snap_to_grid = False
    run = p.add_run()
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-2" \\h \\z \\u'
    run._element.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._element.append(fldChar2)
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar3)
    doc.add_page_break()

def set_update_fields(doc):
    try:
        settings = doc.settings.element
        updateFields = OxmlElement('w:updateFields')
        updateFields.set(qn('w:val'), 'true')
        settings.append(updateFields)
    except Exception:
        pass

def set_outline_level(p, level):
    pPr = p._element.get_or_add_pPr()
    outlineLvl = pPr.find(qn('w:outlineLvl'))
    if outlineLvl is None:
        outlineLvl = OxmlElement('w:outlineLvl')
        inserted = False
        for tag in ['w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr', 'w:pPrChange']:
            child = pPr.find(qn(tag))
            if child is not None:
                child.addprevious(outlineLvl)
                inserted = True
                break
        if not inserted:
            pPr.append(outlineLvl)
    outlineLvl.set(qn('w:val'), str(level - 1))

def is_csv_block(lines):
    if not lines: return False
    if any(any(c in line for c in BORDER_CHARS) for line in lines):
        return False
    return any(',' in line for line in lines if line.strip())

def apply_format_sync(p, font_name, size_pt=10.5, bold=False, base_ind=0.0, hanging_ind=0.0, align_center=False):
    pf = p.paragraph_format
    pf.snap_to_grid = False 
    pf.space_before = Pt(0) 
    pf.space_after = Pt(0)  
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE 
    
    total_left = base_ind + hanging_ind
    pf.left_indent = Pt(10.5 * total_left)
    
    if hanging_ind != 0:
        pf.first_line_indent = Pt(-10.5 * hanging_ind)
    else:
        pf.first_line_indent = Pt(0)
    
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    if not p.runs:
        return
    for run in p.runs:
        set_font(run, font_name, size_pt, bold)

def create_word_doc(text, selected_font, width_mode, add_space_shou, add_space_jou, add_space_kou, add_space_gou, add_space_lv5, add_space_lv6, indent_body_shou, indent_body_jou, indent_body_kou, indent_body_gou, indent_body_lv5, indent_body_lv6, art_indent_val, indent_kou, hanging_kou, indent_gou, hanging_gou, indent_lv5, hanging_lv5, indent_lv6, hanging_lv6, custom_kou, custom_gou, custom_lv5, custom_lv6, out_mode):
    try:
        doc = docx.Document('template.docx')
    except Exception:
        doc = docx.Document()
        
    if "アウトライン" in out_mode:
        set_update_fields(doc)
    
    style = doc.styles['Normal']
    font = style.font
    font.name = selected_font
    font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), selected_font)
    style.paragraph_format.space_after = Pt(0)

    toc_added = False
    art_buffer = []
    csv_buffer = []
    block_buffer = []
    in_code_block = False
    in_csv_table = False
    is_first_text_line = True 
    last_ctx = "none"
    
    def flush_art(lines_to_flush):
        if not lines_to_flush: return
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        tbl_pr = table._element.xpath('w:tblPr')
        if tbl_pr:
            tbl_ind = OxmlElement('w:tblInd')
            indent_twips = int(10.5 * art_indent_val * 20) 
            tbl_ind.set(qn('w:w'), str(indent_twips))
            tbl_ind.set(qn('w:type'), 'dxa')
            tbl_pr[0].append(tbl_ind)
            
        cell = table.cell(0, 0)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(11)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.snap_to_grid = False
        p.paragraph_format.left_indent = Pt(0) 
        
        for i, line in enumerate(lines_to_flush):
            if i > 0: p.add_run('\n')
            converted_line = convert_char_width(line, "全角に統一（推奨）")
            run = p.add_run(converted_line)
            for t in run._element.findall('.//w:t', namespaces=run._element.nsmap):
                t.set(qn('xml:space'), 'preserve')
            set_font(run, 'ＭＳ ゴシック', size_pt=8.5)
        lines_to_flush.clear()

    def flush_csv(lines_to_flush):
        if not lines_to_flush: return
        valid_lines = [l for l in lines_to_flush if l.strip()]
        if not valid_lines: 
            lines_to_flush.clear()
            return
            
        max_cols = max(len(row.split(',')) for row in valid_lines)
        if max_cols <= 1:
            for row in valid_lines:
                p = doc.add_paragraph()
                p.paragraph_format.snap_to_grid = False
                p.paragraph_format.left_indent = Pt(10.5 * indent_body_jou)
                run = p.add_run(convert_char_width(row.strip(), width_mode))
                set_font(run, selected_font, size_pt=10.5)
            lines_to_flush.clear()
            return

        table = doc.add_table(rows=len(valid_lines), cols=max_cols)
        table.style = 'Table Grid'
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r_idx, row_text in enumerate(valid_lines):
            cells_data = row_text.split(',')
            for c_idx, cell_data in enumerate(cells_data):
                if c_idx < len(table.columns):
                    cell = table.cell(r_idx, c_idx)
                    cell.text = cell_data.strip()
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.snap_to_grid = False
                        for run in p.runs:
                            run.text = convert_char_width(run.text, width_mode)
                            set_font(run, selected_font, size_pt=9)
        lines_to_flush.clear()

    def flush_code_block():
        if not block_buffer: return
        if is_csv_block(block_buffer):
            flush_csv(block_buffer)
        else:
            flush_art(block_buffer)

    lines = text.split('\n')
    num_lines = len(lines)

    for i in range(num_lines):
        line = lines[i]
        raw_line = line.rstrip('\r\n')
        raw_line = re.sub(r'\[ \t]?', '', raw_line)
        raw_line = re.sub(r'\[\d+\][ \t]?', '', raw_line)
        line_strip = raw_line.strip()
        
        is_heading = bool(re.match(RE_SHOU, line_strip) or re.match(RE_JOU, line_strip)) or line_strip in TITLE_LIST
        
        if '```' in line_strip or '｀｀｀' in line_strip:
            if in_code_block:
                flush_code_block()
                in_code_block = False
            else:
                if csv_buffer: flush_csv(csv_buffer); in_csv_table = False
                if art_buffer: flush_art(art_buffer)
                in_code_block = True
            continue

        if in_code_block:
            if is_heading:
                flush_code_block()
                in_code_block = False
            else:
                if "The following table:" in line_strip: continue
                if line_strip.lower() in ['text', 'csv', 'markdown']: continue
                block_buffer.append(raw_line)
                continue

        if not line_strip:
            if csv_buffer: flush_csv(csv_buffer); in_csv_table = False
            if art_buffer: flush_art(art_buffer)
            continue

        if "The following table:" in line_strip:
            if art_buffer: flush_art(art_buffer)
            in_csv_table = True
            continue
            
        if in_csv_table:
            if is_heading:
                flush_csv(csv_buffer)
                in_csv_table = False
            else:
                csv_buffer.append(raw_line)
                continue

        if any(c in raw_line for c in BORDER_CHARS):
            art_buffer.append(raw_line)
            continue
        else:
            if art_buffer: flush_art(art_buffer)

        # カッコ見出しの先読み判定（空行スキップ対応）
        line_strip_conv = convert_char_width(line_strip, width_mode)
        is_midashi_go = bool(re.match(RE_MIDASHI_GO, line_strip_conv))
        next_is_jou = False
        if is_midashi_go:
            for j in range(i + 1, num_lines):
                nxt = lines[j].strip()
                if not nxt: continue # 空行はスキップして次を見る
                if re.match(RE_JOU, convert_char_width(nxt, width_mode)):
                    next_is_jou = True
                break

        line_strip = convert_char_width(line_strip, width_mode)
        
        m_shou = get_hierarchy_match(line_strip, RE_SHOU, None, width_mode)
        m_jou = get_hierarchy_match(line_strip, RE_JOU, None, width_mode)
        m_kou = get_hierarchy_match(line_strip, RE_KOU, custom_kou, width_mode)
        m_gou = get_hierarchy_match(line_strip, RE_GOU, custom_gou, width_mode)
        m_lv5 = get_hierarchy_match(line_strip, RE_LV5, custom_lv5, width_mode)
        m_lv6 = get_hierarchy_match(line_strip, RE_LV6, custom_lv6, width_mode)

        if m_kou: line_strip = normalize_marker_space(line_strip, m_kou)
        elif m_gou: line_strip = normalize_marker_space(line_strip, m_gou)
        elif m_lv5: line_strip = normalize_marker_space(line_strip, m_lv5)
        elif m_lv6: line_strip = normalize_marker_space(line_strip, m_lv6)

        is_title = False
        if is_first_text_line:
            if not m_shou and not m_jou:
                is_title = True
            is_first_text_line = False
        elif line_strip in TITLE_LIST:
            is_title = True

        if (m_shou or m_jou) and not toc_added:
            add_toc(doc)
            toc_added = True

        if is_title:
            p = doc.add_paragraph()
            run = p.add_run(line_strip)
            apply_format_sync(p, selected_font, size_pt=13.5, bold=True, align_center=True)
            p.paragraph_format.space_after = Pt(24)
            last_ctx = "shou"

        # 条文の上の見出し語の場合
        elif is_midashi_go and next_is_jou:
            p = doc.add_paragraph(line_strip)
            apply_format_sync(p, selected_font, size_pt=10.5, bold=True, base_ind=0.0)
            p.paragraph_format.space_before = Pt(12) 
            last_ctx = "jou"

        elif m_shou:
            p = doc.add_paragraph(line_strip)
            if "アウトライン" in out_mode: set_outline_level(p, 1)
            apply_format_sync(p, selected_font, size_pt=13, bold=True, align_center=True)
            p.paragraph_format.space_before = Pt(12) 
            last_ctx = "shou"
            
        elif m_jou:
            p = doc.add_paragraph(line_strip)
            if "アウトライン" in out_mode: set_outline_level(p, 2)
            apply_format_sync(p, selected_font, size_pt=10.5, bold=True, base_ind=0.0)
            p.paragraph_format.space_before = Pt(6) 
            last_ctx = "jou"
        
        elif m_kou:
            p = doc.add_paragraph(line_strip)
            if "アウトライン" in out_mode: set_outline_level(p, 3)
            apply_format_sync(p, selected_font, size_pt=10.5, bold=False, base_ind=indent_kou, hanging_ind=hanging_kou)
            last_ctx = "kou"

        elif m_gou:
            p = doc.add_paragraph(line_strip)
            if "アウトライン" in out_mode: set_outline_level(p, 4)
            apply_format_sync(p, selected_font, size_pt=10.5, bold=False, base_ind=indent_gou, hanging_ind=hanging_gou)
            last_ctx = "gou"

        elif m_lv5:
            p = doc.add_paragraph(line_strip)
            if "アウトライン" in out_mode: set_outline_level(p, 5)
            apply_format_sync(p, selected_font, size_pt=10.5, bold=False, base_ind=indent_lv5, hanging_ind=hanging_lv5)
            last_ctx = "lv5"

        elif m_lv6:
            p = doc.add_paragraph(line_strip)
            if "アウトライン" in out_mode: set_outline_level(p, 6)
            apply_format_sync(p, selected_font, size_pt=10.5, bold=False, base_ind=indent_lv6, hanging_ind=hanging_lv6)
            last_ctx = "lv6"

        else:
            p = doc.add_paragraph()
            
            add_space = False
            if last_ctx in ["shou", "title", "none"]:
                ind = indent_body_shou
                add_space = add_space_shou
            elif last_ctx == "jou":
                ind = indent_body_jou
                add_space = add_space_jou
            elif last_ctx == "kou":
                ind = indent_body_kou
                add_space = add_space_kou
            elif last_ctx == "gou":
                ind = indent_body_gou
                add_space = add_space_gou
            elif last_ctx == "lv5":
                ind = indent_body_lv5
                add_space = add_space_lv5
            elif last_ctx == "lv6":
                ind = indent_body_lv6
                add_space = add_space_lv6
            
            output_text = line_strip
            if add_space and output_text:
                output_text = '　' + output_text
            run = p.add_run(output_text)
            
            apply_format_sync(p, selected_font, size_pt=10.5, bold=False, base_ind=ind, hanging_ind=0.0)
            
    if block_buffer: flush_code_block()
    if art_buffer: flush_art(art_buffer)
    if csv_buffer: flush_csv(csv_buffer)
        
    if not toc_added: add_toc(doc)
    return doc

def get_html_table(csv_lines, width_mode):
    csv_lines = [l for l in csv_lines if l.strip()]
    if not csv_lines: return ""
    max_cols = max(len(row.split(',')) for row in csv_lines)
    if max_cols <= 1:
        return "".join([f"<div style='margin-left: 1.0em;'>{convert_char_width(row.strip(), width_mode)}</div>" for row in csv_lines])

    html = "<table style='width:100%; border-collapse: collapse; margin: 10px 0; font-size: 12px;'>"
    for row in csv_lines:
        html += "<tr>"
        for cell in row.split(','):
            html += f"<td style='border: 1px solid #999; padding: 4px 8px; text-align: center;'>{convert_char_width(cell.strip(), width_mode)}</td>"
        html += "</tr>"
    html += "</table>"
    return html

def get_html_art(art_lines, indent_art, width_mode):
    if not art_lines: return ""
    converted_lines = [convert_char_width(line, width_mode) for line in art_lines]
    html = f"<div style='background-color: #fff3cd; color: #856404; padding: 4px 8px; font-size: 11px; margin-top: 10px; border-radius: 3px; margin-left: {indent_art}em; width: fit-content;'>"
    html += "⚠️ 本ツールでは可能な限り自動整形を行っておりますが、出力後はユーザー側で適宜微調整をお願いいたします。</div>"
    html += f"<pre style='border: 1px solid #999; background-color: #fff; padding: 10px; font-family: \"MS Gothic\", \"ＭＳ ゴシック\", monospace; font-size: 13px; line-height: 1.2; margin: 0 0 10px; margin-left: {indent_art}em; white-space: pre; overflow-x: auto;'>{chr(10).join(converted_lines)}</pre>"
    return html

def reset_settings():
    st.session_state.font_choice = "MS 明朝"
    st.session_state.width_mode_choice = "全角に統一（推奨）"
    st.session_state.out_mode = "【推奨】アウトライン連動"
    
    st.session_state.as_shou = False
    st.session_state.as_jou = False
    st.session_state.as_kou = False
    st.session_state.as_gou = False
    st.session_state.as_l5 = False
    st.session_state.as_l6 = False
    
    st.session_state.in_body_shou = 0.0
    st.session_state.in_body_jou = 1.0
    st.session_state.in_body_kou = 4.0
    st.session_state.in_body_gou = 5.0
    st.session_state.in_body_lv5 = 7.0
    st.session_state.in_body_lv6 = 9.0
    
    st.session_state.ik = 2.0
    st.session_state.hk = 2.0
    st.session_state.ig = 4.0
    st.session_state.hg = 1.0
    st.session_state.i5 = 4.0
    st.session_state.h5 = 3.0
    st.session_state.i6 = 6.0
    st.session_state.h6 = 3.0
    
    st.session_state.art_type = "本文（番号なし）の開始位置に合わせる"
    st.session_state.c_kou = ""
    st.session_state.c_gou = ""
    st.session_state.c_lv5 = ""
    st.session_state.c_lv6 = ""
    st.session_state.auto_renumber = False

def main():
    st.set_page_config(page_title="就業規則 Word変換", layout="wide")
    st.title("📄 就業規則プロフェッショナル整形ツール")
    
    if "initialized" not in st.session_state:
        reset_settings()
        st.session_state.initialized = True

    st.sidebar.button("🔄 推奨設定（初期値）に戻す（上）", on_click=reset_settings, key="reset_top", use_container_width=True)

    with st.sidebar.form(key='settings_form'):
        submit_button_top = st.form_submit_button(label='🔄 設定を適用してプレビュー更新（上）', type='primary', use_container_width=True)

        st.header("🚀 出力エンジン・書式設定")
        out_mode = st.radio("モード選択", ["【推奨】アウトライン連動", "直接モード"], key="out_mode")
        if "アウトライン" in out_mode:
            df_outline = pd.DataFrame({"Lv": ["1", "2", "3", "4", "5", "6"], "役割": ["章", "条", "項", "号", "(1)", "(ア)"], "内部処理": ["アウトラインLv1", "アウトラインLv2", "アウトラインLv3", "アウトラインLv4", "アウトラインLv5", "アウトラインLv6"]})
            st.table(df_outline)

        st.markdown("---")
        st.write("**🔢 自動採番オプション**")
        auto_renumber = st.checkbox("章・条の番号を自動で振り直す（1から連番）", key="auto_renumber")
        if auto_renumber:
            st.warning("⚠️ 見出しの番号のみ更新されます。本文中の参照部分は手動で修正してください。")

        st.markdown("---")
        selected_font = st.selectbox("基本フォント", options=["MS 明朝", "MS ゴシック", "Meiryo", "Yu Mincho", "Yu Gothic", "Arial"], key="font_choice")
        width_mode = st.radio("英数字の表記統一", options=["全角に統一（推奨）", "半角に統一", "変換しない（元のまま）"], key="width_mode_choice")
        
        st.markdown("---")
        st.write("**【第1階層】章（表題）の設定**")
        add_space_shou = st.checkbox("本文の先頭に全角スペースを自動挿入", key='as_shou')
        indent_body_shou = st.slider("章 直後の本文 字下げ", min_value=0.0, max_value=10.0, step=0.5, key="in_body_shou")
        
        st.markdown("---")
        st.write("**【第2階層】条の設定**")
        add_space_jou = st.checkbox("本文の先頭に全角スペースを自動挿入", key='as_jou')
        indent_body_jou = st.slider("条 直後の本文 字下げ", min_value=0.0, max_value=10.0, step=0.5, key="in_body_jou")

        st.markdown("---")
        st.write("**【第3階層】項（1. 等）の設定**")
        indent_kou = st.slider("項 の字下げ", min_value=0.0, max_value=15.0, step=0.5, key="ik")
        hanging_kou = st.slider("項 の突き出し幅", min_value=0.0, max_value=5.0, step=0.5, key="hk")
        add_space_kou = st.checkbox("本文の先頭に全角スペースを自動挿入", key='as_kou')
        indent_body_kou = st.slider("項 直後の本文 字下げ", min_value=0.0, max_value=15.0, step=0.5, key="in_body_kou")

        st.markdown("---")
        st.write("**【第4階層】号（① 等）の設定**")
        indent_gou = st.slider("号 の字下げ", min_value=0.0, max_value=15.0, step=0.5, key="ig")
        hanging_gou = st.slider("号 の突き出し幅", min_value=0.0, max_value=5.0, step=0.5, key="hg")
        add_space_gou = st.checkbox("本文の先頭に全角スペースを自動挿入", key='as_gou')
        indent_body_gou = st.slider("号 直後の本文 字下げ", min_value=0.0, max_value=15.0, step=0.5, key="in_body_gou")

        st.markdown("---")
        st.write("**【第5階層】(1) 等の設定**")
        indent_lv5 = st.slider("第5階層 の字下げ", min_value=0.0, max_value=15.0, step=0.5, key="i5")
        hanging_lv5 = st.slider("第5階層 の突き出し幅", min_value=0.0, max_value=5.0, step=0.5, key="h5")
        add_space_lv5 = st.checkbox("本文の先頭に全角スペースを自動挿入", key='as_l5')
        indent_body_lv5 = st.slider("第5階層 直後の本文 字下げ", min_value=0.0, max_value=15.0, step=0.5, key="in_body_lv5")

        st.markdown("---")
        st.write("**【第6階層】(ア) 等の設定**")
        indent_lv6 = st.slider("第6階層 の字下げ", min_value=0.0, max_value=15.0, step=0.5, key="i6")
        hanging_lv6 = st.slider("第6階層 の突き出し幅", min_value=0.0, max_value=5.0, step=0.5, key="h6")
        add_space_lv6 = st.checkbox("本文の先頭に全角スペースを自動挿入", key='as_l6')
        indent_body_lv6 = st.slider("第6階層 直後の本文 字下げ", min_value=0.0, max_value=15.0, step=0.5, key="in_body_lv6")

        st.markdown("---")
        with st.expander("⚙️ 特殊な記号を追加（上級者向け）"):
            st.write("自社の規定で特殊な記号を使っている場合のみ入力。（複数ある場合はカンマ区切り）")
            custom_kou = st.text_input("第3階層（項）に追加", placeholder="例: ◆, ●", key="c_kou")
            custom_gou = st.text_input("第4階層（号）に追加", placeholder="例: ・", key="c_gou")
            custom_lv5 = st.text_input("第5階層に追加", placeholder="例: a.", key="c_lv5")
            custom_lv6 = st.text_input("第6階層に追加", placeholder="例: (a)", key="c_lv6")

        st.markdown("---")
        st.write("**表（ツリー図）の配置設定**")
        art_indent_type = st.radio(
            "表の左端をどこに合わせますか？",
            options=["左端（0字下げ）に合わせる", "本文（番号なし）の開始位置に合わせる", "項（１．等）の番号位置に合わせる", "号（① 等）の番号位置に合わせる"],
            key="art_type"
        )
        
        st.markdown("---")
        submit_button_bottom = st.form_submit_button(label='🔄 設定を適用してプレビュー更新（下）', type='primary', use_container_width=True)

    st.sidebar.button("🔄 推奨設定（初期値）に戻す（下）", on_click=reset_settings, key="reset_bottom", use_container_width=True)

    calculated_art_indent = 0.0
    if art_indent_type == "本文（番号なし）の開始位置に合わせる":
        calculated_art_indent = indent_body_jou
    elif art_indent_type == "項（１．等）の番号位置に合わせる":
        calculated_art_indent = indent_kou
    elif art_indent_type == "号（① 等）の番号位置に合わせる":
        calculated_art_indent = indent_gou

    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.write("▼ ここにテキストを貼り付け（貼り付け後、左の「設定を適用」ボタンを押してください）")
        text_input = st.text_area("就業規則テキスト", height=600, label_visibility="collapsed")
    
    with col2:
        st.write("▼ インデント簡易プレビュー（全行表示）")
        
        processed_text = text_input
        change_log = []
        if auto_renumber and processed_text:
            processed_text, change_log = renumber_headings(processed_text)
            
        if change_log:
            st.success("🔄 条文番号を自動で振り直しました。本文中の参照番号（「第〇条の規定により～」など）を修正する際の参考にしてください。")
            st.dataframe(pd.DataFrame(change_log), use_container_width=True)

        if processed_text:
            preview_html = f"<div style='font-family: \"{selected_font}\", sans-serif; font-size: 14px; line-height: 1.2; border: 1px solid #ddd; padding: 10px; height: 600px; overflow-y: auto;'>"
            
            art_buffer = []
            csv_buffer = []
            block_buffer = []
            in_code_block = False
            in_csv_table = False
            is_first_text_line = True 
            last_p = "none"
            
            lines_preview = processed_text.split('\n')
            num_lines_preview = len(lines_preview)

            for i in range(num_lines_preview):
                line = lines_preview[i]
                raw_line = line.rstrip('\r\n')
                raw_line = re.sub(r'\[ \t]?', '', raw_line)
                raw_line = re.sub(r'\[\d+\][ \t]?', '', raw_line)
                line_strip = raw_line.strip()
                
                is_heading = bool(re.match(RE_SHOU, line_strip) or re.match(RE_JOU, line_strip)) or line_strip in TITLE_LIST
                
                if '```' in line_strip or '｀｀｀' in line_strip:
                    if in_code_block:
                        if is_csv_block(block_buffer): preview_html += get_html_table(block_buffer, width_mode)
                        else: preview_html += get_html_art(block_buffer, calculated_art_indent, width_mode)
                        block_buffer.clear()
                        in_code_block = False
                    else:
                        if csv_buffer: preview_html += get_html_table(csv_buffer, width_mode); csv_buffer.clear(); in_csv_table = False
                        if art_buffer: preview_html += get_html_art(art_buffer, calculated_art_indent, width_mode); art_buffer.clear()
                        in_code_block = True
                    continue

                if in_code_block:
                    if is_heading:
                        if is_csv_block(block_buffer): preview_html += get_html_table(block_buffer, width_mode)
                        else: preview_html += get_html_art(block_buffer, calculated_art_indent, width_mode)
                        block_buffer.clear()
                        in_code_block = False
                    else:
                        if "The following table:" in line_strip: continue
                        if line_strip.lower() in ['text', 'csv', 'markdown']: continue
                        block_buffer.append(raw_line)
                        continue

                if not line_strip:
                    if csv_buffer: preview_html += get_html_table(csv_buffer, width_mode); csv_buffer.clear(); in_csv_table = False
                    if art_buffer: preview_html += get_html_art(art_buffer, calculated_art_indent, width_mode); art_buffer.clear()
                    continue

                if "The following table:" in line_strip:
                    if art_buffer: preview_html += get_html_art(art_buffer, calculated_art_indent, width_mode); art_buffer.clear()
                    in_csv_table = True
                    continue
                    
                if in_csv_table:
                    if is_heading:
                        preview_html += get_html_table(csv_buffer, width_mode); csv_buffer.clear(); in_csv_table = False
                    else:
                        csv_buffer.append(raw_line)
                        continue

                if any(c in raw_line for c in BORDER_CHARS):
                    art_buffer.append(raw_line)
                    continue
                else:
                    if art_buffer: 
                        preview_html += get_html_art(art_buffer, calculated_art_indent, width_mode)
                        art_buffer.clear()
                
                # HTMLプレビュー側での見出し先読み
                line_strip_conv = convert_char_width(line_strip, width_mode)
                is_midashi_go = bool(re.match(RE_MIDASHI_GO, line_strip_conv))
                next_is_jou = False
                if is_midashi_go:
                    for j in range(i + 1, num_lines_preview):
                        nxt = lines_preview[j].strip()
                        if not nxt: continue
                        if re.match(RE_JOU, convert_char_width(nxt, width_mode)):
                            next_is_jou = True
                        break

                line_strip = convert_char_width(line_strip, width_mode)
                
                m_sh = get_hierarchy_match(line_strip, RE_SHOU, None, width_mode)
                m_jo = get_hierarchy_match(line_strip, RE_JOU, None, width_mode)
                m_ko = get_hierarchy_match(line_strip, RE_KOU, custom_kou, width_mode)
                m_go = get_hierarchy_match(line_strip, RE_GOU, custom_gou, width_mode)
                m_l5 = get_hierarchy_match(line_strip, RE_LV5, custom_lv5, width_mode)
                m_l6 = get_hierarchy_match(line_strip, RE_LV6, custom_lv6, width_mode)

                if m_ko: line_strip = normalize_marker_space(line_strip, m_ko)
                elif m_go: line_strip = normalize_marker_space(line_strip, m_go)
                elif m_l5: line_strip = normalize_marker_space(line_strip, m_l5)
                elif m_l6: line_strip = normalize_marker_space(line_strip, m_l6)

                is_title = False
                if is_first_text_line:
                    if not m_sh and not m_jo:
                        is_title = True
                    is_first_text_line = False
                elif line_strip in TITLE_LIST:
                    is_title = True

                if is_title:
                    preview_html += f"<div style='text-align: center; font-weight: bold; margin-bottom: 24px; font-size: 1.3em;'>{line_strip}</div>"
                    last_p = "title"

                elif is_midashi_go and next_is_jou:
                    preview_html += f"<div style='font-weight: bold; margin-top: 15px;'>{line_strip}</div>"
                    last_p = "jou"

                elif m_sh:
                    preview_html += f"<div style='text-align: center; font-weight: bold; margin-top: 5px;'>{line_strip}</div>"; last_p = "shou"
                elif m_jo:
                    preview_html += f"<div style='font-weight: bold; margin-top: 5px;'>{line_strip}</div>"; last_p = "jou"
                
                elif m_ko:
                    preview_html += f"<div style='margin-left: {indent_kou}em; padding-left: {hanging_kou}em; text-indent: -{hanging_kou}em; color: red;'>{line_strip}</div>"; last_p = "kou"
                elif m_go:
                    preview_html += f"<div style='margin-left: {indent_gou}em; padding-left: {hanging_gou}em; text-indent: -{hanging_gou}em; color: red;'>{line_strip}</div>"; last_p = "gou"
                elif m_l5:
                    preview_html += f"<div style='margin-left: {indent_lv5}em; padding-left: {hanging_lv5}em; text-indent: -{hanging_lv5}em; color: red;'>{line_strip}</div>"; last_p = "lv5"
                elif m_l6:
                    preview_html += f"<div style='margin-left: {indent_lv6}em; padding-left: {hanging_lv6}em; text-indent: -{hanging_lv6}em; color: red;'>{line_strip}</div>"; last_p = "lv6"
                
                else:
                    add_space = False
                    if last_p in ["shou", "title", "none"]:
                        ind = indent_body_shou
                        add_space = add_space_shou
                    elif last_p == "jou":
                        ind = indent_body_jou
                        add_space = add_space_jou
                    elif last_p == "kou":
                        ind = indent_body_kou
                        add_space = add_space_kou
                    elif last_p == "gou":
                        ind = indent_body_gou
                        add_space = add_space_gou
                    elif last_p == "lv5":
                        ind = indent_body_lv5
                        add_space = add_space_lv5
                    elif last_p == "lv6":
                        ind = indent_body_lv6
                        add_space = add_space_lv6
                    
                    output_text = line_strip
                    if add_space and output_text:
                        output_text = '　' + output_text
                    
                    preview_html += f"<div style='margin-left: {ind}em; color:#555;'>{output_text}</div>"
            
            if block_buffer:
                if is_csv_block(block_buffer): preview_html += get_html_table(block_buffer, width_mode)
                else: preview_html += get_html_art(block_buffer, calculated_art_indent, width_mode)
            if csv_buffer: preview_html += get_html_table(csv_buffer, width_mode)
            if art_buffer: preview_html += get_html_art(art_buffer, calculated_art_indent, width_mode)
                
            preview_html += "</div>"
            st.markdown(preview_html, unsafe_allow_html=True)
        else:
            st.info("テキストを入力すると、ここにインデントのプレビューが表示されます。")

    if text_input:
        st.markdown("---")
        doc = create_word_doc(processed_text, selected_font, width_mode, add_space_shou, add_space_jou, add_space_kou, add_space_gou, add_space_lv5, add_space_lv6, indent_body_shou, indent_body_jou, indent_body_kou, indent_body_gou, indent_body_lv5, indent_body_lv6, calculated_art_indent, indent_kou, hanging_kou, indent_gou, hanging_gou, indent_lv5, hanging_lv5, indent_lv6, hanging_lv6, custom_kou, custom_gou, custom_lv5, custom_lv6, out_mode)
        bio = io.BytesIO()
        doc.save(bio)
        
        st.download_button(
            label=f"📥 {out_mode}設定を反映してWordファイルをダウンロード",
            data=bio.getvalue(),
            file_name="整形済み就業規則.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

if __name__ == "__main__":
    main()
